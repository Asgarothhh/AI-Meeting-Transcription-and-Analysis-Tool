import os
import sys
import markdown
import torch
import time
from pathlib import Path

from PyQt6.QtGui import QPageLayout
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout,
                             QPushButton, QLabel, QFileDialog, QProgressBar,
                             QTextBrowser, QHBoxLayout, QCheckBox, QMessageBox,
                             QComboBox)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QMarginsF
from PyQt6.QtPrintSupport import QPrinter

from src.utils import process_media
from src.transcription import TranscriptionPipeline
from src.summarization import graph
from src.config import Config
from langchain_text_splitters import RecursiveCharacterTextSplitter
from docx import Document
import whisper
from dotenv import load_dotenv

load_dotenv()

config = Config()
DEVICE = config["device"] or ("cuda" if torch.cuda.is_available() else "cpu")

PROJECT_ROOT = Path(__file__).resolve().parent.parent
TEMP_DIR = PROJECT_ROOT / config["paths"]["temp_dir"]
TEMP_DIR.mkdir(parents=True, exist_ok=True)

PROTON_STYLESHEET = """
QMainWindow, QWidget {
    background-color: #14141c;
    color: #f1f1f4;
    font-family: "Inter", "Segoe UI", sans-serif;
}
QPushButton {
    background-color: #6d4aff;
    color: white;
    border: none;
    border-radius: 6px;
    padding: 8px 16px;
    font-weight: bold;
    font-size: 13px;
}
QPushButton:hover {
    background-color: #8c72ff;
}
QPushButton:disabled {
    background-color: #313145;
    color: #6c6c85;
}
QProgressBar {
    background-color: #1f1f2e;
    border: 1px solid #313145;
    border-radius: 6px;
    text-align: center;
    color: white;
}
QProgressBar::chunk {
    background-color: #6d4aff;
    border-radius: 5px;
}
QTextBrowser {
    background-color: #1f1f2e;
    border: 1px solid #313145;
    border-radius: 8px;
    padding: 15px;
    font-size: 14px;
    line-height: 1.5;
}
QCheckBox { font-size: 13px; }
QCheckBox::indicator {
    width: 18px; height: 18px; border-radius: 4px;
    border: 2px solid #6d4aff; background-color: #1f1f2e;
}
QCheckBox::indicator:checked { background-color: #6d4aff; }
QLabel { font-size: 13px; }
QComboBox {
    background-color: #1f1f2e;
    border: 1px solid #313145;
    border-radius: 6px;
    padding: 5px 10px;
    color: white;
    font-size: 13px;
}
QComboBox::drop-down { border: none; }
"""


class ProcessingThread(QThread):
    progress = pyqtSignal(int, str)
    finished = pyqtSignal(str, str)
    error = pyqtSignal(str)

    def __init__(self, file_path, apply_noise_reduction, language_code):
        super().__init__()
        self.file_path = file_path
        self.apply_noise_reduction = apply_noise_reduction
        self.language_code = language_code

    @staticmethod
    def _speaker_from_segment(seg: dict) -> str:
        speaker = seg.get("speaker")
        if speaker:
            return speaker

        words = seg.get("words") or []
        word_speakers = [w.get("speaker") for w in words if w.get("speaker")]
        if word_speakers:
            return max(set(word_speakers), key=word_speakers.count)
        return "SPEAKER_UNKNOWN"

    def _segments_to_speaker_turns(self, segments: list[dict]) -> list[dict]:
        """
        Преобразует сегменты WhisperX в более реалистичные реплики спикеров.
        Приоритет: word-level speaker labels; fallback: speaker сегмента.
        """
        turns: list[dict] = []

        for seg in segments:
            words = seg.get("words") or []
            if not words:
                text = seg.get("text", "").strip()
                if text:
                    turns.append(
                        {
                            "speaker": self._speaker_from_segment(seg),
                            "start": float(seg.get("start", 0)),
                            "end": float(seg.get("end", 0)),
                            "text": text,
                        }
                    )
                continue

            current_turn = None
            segment_fallback_speaker = self._speaker_from_segment(seg)

            for word in words:
                token = (word.get("word") or "").strip()
                if not token:
                    continue

                speaker = word.get("speaker") or segment_fallback_speaker
                start = float(word.get("start", seg.get("start", 0)))
                end = float(word.get("end", start))

                if current_turn is None or current_turn["speaker"] != speaker:
                    if current_turn is not None:
                        turns.append(current_turn)
                    current_turn = {"speaker": speaker, "start": start, "end": end, "words": [token]}
                else:
                    current_turn["end"] = end
                    current_turn["words"].append(token)

            if current_turn is not None:
                turns.append(current_turn)

        # Нормализуем текст реплик и убираем пустые.
        normalized_turns: list[dict] = []
        for turn in turns:
            if "words" in turn:
                text = " ".join(turn["words"]).strip()
            else:
                text = (turn.get("text") or "").strip()
            if not text:
                continue
            normalized_turns.append(
                {
                    "speaker": turn["speaker"],
                    "start": float(turn["start"]),
                    "end": float(turn["end"]),
                    "text": text,
                }
            )

        return normalized_turns

    @staticmethod
    def _merge_fragmented_turns(turns: list[dict]) -> list[dict]:
        """Склеивает слишком короткие фрагменты, которые часто появляются на стыке слов."""
        if not turns:
            return turns

        merged: list[dict] = []
        i = 0
        while i < len(turns):
            current = dict(turns[i])

            # Если текущая реплика очень короткая, это обычно хвост фразы.
            duration = current["end"] - current["start"]
            if duration <= 0.6 and i > 0:
                prev = merged[-1] if merged else None
                nxt = turns[i + 1] if i + 1 < len(turns) else None

                if prev and prev["speaker"] == current["speaker"]:
                    prev["end"] = max(prev["end"], current["end"])
                    prev["text"] = f"{prev['text']} {current['text']}".strip()
                    i += 1
                    continue

                # Если короткий фрагмент зажат между одинаковыми спикерами — склеиваем.
                if prev and nxt and prev["speaker"] == nxt["speaker"]:
                    prev["end"] = max(prev["end"], nxt["end"])
                    prev["text"] = f"{prev['text']} {current['text']} {nxt['text']}".strip()
                    i += 2
                    continue

            merged.append(current)
            i += 1

        return merged

    @staticmethod
    def _rebalance_two_speaker_runs(turns: list[dict]) -> list[dict]:
        """
        Если в 2-спикерном диалоге есть длинная серия одного спикера,
        мягко переразмечает середину серии, чтобы не залипать в один label.
        """
        if not turns:
            return turns

        speakers = sorted({t["speaker"] for t in turns if t.get("speaker") and t["speaker"] != "SPEAKER_UNKNOWN"})
        if len(speakers) != 2:
            return turns

        adjusted = [dict(t) for t in turns]
        i = 0
        while i < len(adjusted):
            j = i
            while j + 1 < len(adjusted) and adjusted[j + 1]["speaker"] == adjusted[i]["speaker"]:
                j += 1

            run_len = j - i + 1
            if run_len >= 3:
                run_duration = sum(max(0.0, adjusted[k]["end"] - adjusted[k]["start"]) for k in range(i, j + 1))
                # Применяем только для заметно длинных серий.
                if run_duration >= 5.0:
                    other = speakers[0] if adjusted[i]["speaker"] == speakers[1] else speakers[1]
                    for k in range(i + 1, j + 1, 2):
                        # Сохраняем первую реплику серии исходной, а далее мягко чередуем.
                        adjusted[k]["speaker"] = other

            i = j + 1

        return adjusted

    def run(self):
        timestamp = int(time.time())
        temp_audio_path = str(TEMP_DIR / f"temp_proc_{timestamp}.wav")
        try:
            self.progress.emit(10, "Подготовка медиафайла...")
            process_media(self.file_path, temp_audio_path, self.apply_noise_reduction)

            self.progress.emit(30, "Транскрибация и диаризация...")
            pipeline = TranscriptionPipeline()
            pipeline.language = self.language_code

            result = pipeline.run(temp_audio_path)
            segments = result.get("segments", [])
            turns = self._segments_to_speaker_turns(segments)
            turns = self._merge_fragmented_turns(turns)
            turns = self._rebalance_two_speaker_runs(turns)

            transcript_lines = []
            for turn in turns:
                speaker = turn["speaker"]
                text = turn["text"]
                start = turn["start"]
                end = turn["end"]
                transcript_lines.append(f"**{speaker}** [{start:.1f} - {end:.1f}]: {text}")

            full_transcript_md = "\n\n".join(transcript_lines)
            transcript_string = "\n".join(
                [f"{t['speaker']}: {t['text']}" for t in turns]
            )

            self.progress.emit(70, "Анализ и суммаризация (LLM)...")
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=3500, chunk_overlap=200)
            chunks = text_splitter.split_text(transcript_string)

            summary_result = graph.invoke({"contents": chunks})
            final_summary = summary_result.get("final_summary", "Не удалось сгенерировать саммари.")

            self.progress.emit(100, "Готово!")
            self.finished.emit(final_summary, full_transcript_md)

        except Exception as e:
            self.error.emit(str(e))
        finally:
            if os.path.exists(temp_audio_path):
                try:
                    os.remove(temp_audio_path)
                except Exception as e:
                    print(f"Не удалось удалить файл: {e}")


class ProtonApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Meeting Intelligence")
        self.setMinimumSize(950, 700)
        self.setStyleSheet(PROTON_STYLESHEET)

        self.current_summary = ""
        self.current_transcript = ""
        self.selected_file = None
        self.thread = None

        # Инициализация переменных UI (чтобы избежать предупреждений PEP8)
        self.btn_select = None
        self.lbl_file = None
        self.chk_noise = None
        self.combo_lang = None
        self.btn_start = None
        self.progress_bar = None
        self.lbl_status = None
        self.lbl_device = None
        self.text_viewer = None
        self.btn_export_md = None
        self.btn_export_docx = None
        self.btn_export_pdf = None

        self.init_ui()

    def init_ui(self):
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        layout = QVBoxLayout(main_widget)
        layout.setSpacing(15)
        layout.setContentsMargins(20, 20, 20, 20)

        # Контролы (Верхняя панель)
        control_layout = QHBoxLayout()

        self.btn_select = QPushButton("Выбрать Файл")
        self.btn_select.clicked.connect(self.select_file)

        self.lbl_file = QLabel("Файл не выбран")
        self.lbl_file.setStyleSheet("color: #a0a0b5;")

        # Выбор языка
        self.combo_lang = QComboBox()
        self.combo_lang.addItem("🌐 Автоопределение", None)
        sorted_languages = sorted(whisper.tokenizer.LANGUAGES.items(), key=lambda x: x[1])

        for code, name in sorted_languages:
            self.combo_lang.addItem(f"{name.capitalize()}", code)

        self.chk_noise = QCheckBox("Фильтр шумов")

        self.btn_start = QPushButton("Начать обработку")
        self.btn_start.clicked.connect(self.start_processing)
        self.btn_start.setEnabled(False)

        control_layout.addWidget(self.btn_select)
        control_layout.addWidget(self.lbl_file)
        control_layout.addStretch()
        control_layout.addWidget(self.combo_lang)
        control_layout.addWidget(self.chk_noise)
        control_layout.addWidget(self.btn_start)

        layout.addLayout(control_layout)

        # Статус и Прогресс
        status_layout = QHBoxLayout()
        self.lbl_status = QLabel("Ожидание...")

        # Индикатор устройства
        device_color = "#4CAF50" if DEVICE == "cuda" else "#f44336"
        self.lbl_device = QLabel(
            f"⚙️ Device: <span style='color:{device_color}; font-weight:bold;'>{DEVICE.upper()}</span>")

        status_layout.addWidget(self.lbl_status)
        status_layout.addStretch()
        status_layout.addWidget(self.lbl_device)

        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        self.progress_bar.setTextVisible(False)

        layout.addLayout(status_layout)
        layout.addWidget(self.progress_bar)

        # Просмотрщик результата
        self.text_viewer = QTextBrowser()
        self.text_viewer.setOpenExternalLinks(True)
        layout.addWidget(self.text_viewer)

        # Панель экспорта (Нижняя)
        export_layout = QHBoxLayout()
        self.btn_export_md = QPushButton("Скачать .md")
        self.btn_export_docx = QPushButton("Скачать .docx")
        self.btn_export_pdf = QPushButton("Скачать .pdf")

        for btn in [self.btn_export_md, self.btn_export_docx, self.btn_export_pdf]:
            btn.setEnabled(False)
            export_layout.addWidget(btn)

        self.btn_export_md.clicked.connect(lambda: self.export_file("md"))
        self.btn_export_docx.clicked.connect(lambda: self.export_file("docx"))
        self.btn_export_pdf.clicked.connect(lambda: self.export_file("pdf"))

        export_layout.addStretch()
        layout.addLayout(export_layout)

    def select_file(self):
        file_name, _ = QFileDialog.getOpenFileName(
            self, "Выберите медиафайл", "",
            "Media Files (*.mp3 *.wav *.m4a *.mp4 *.mkv *.avi *.mov)"
        )
        if file_name:
            allowed_exts = {".mp3", ".wav", ".m4a", ".mp4", ".mkv", ".avi", ".mov"}
            suffix = Path(file_name).suffix.lower()
            if suffix not in allowed_exts:
                QMessageBox.warning(self, "Неподдерживаемый формат", f"Формат {suffix} не поддерживается.")
                return

            self.selected_file = file_name
            self.lbl_file.setText(Path(file_name).name)
            self.lbl_file.setToolTip(file_name)
            self.btn_start.setEnabled(True)

    def start_processing(self):
        if not self.selected_file or not Path(self.selected_file).exists():
            QMessageBox.warning(self, "Файл не найден", "Выберите существующий аудио/видеофайл.")
            return

        self.btn_start.setEnabled(False)
        self.btn_select.setEnabled(False)
        self.progress_bar.setValue(0)
        self.text_viewer.clear()

        # Получаем код языка из Combobox
        selected_lang_code = self.combo_lang.currentData()
        noise_reduction = self.chk_noise.isChecked()

        self.thread = ProcessingThread(self.selected_file, noise_reduction, selected_lang_code)
        self.thread.progress.connect(self.update_progress)
        self.thread.finished.connect(self.on_finished)
        self.thread.error.connect(self.on_error)
        self.thread.start()

    def update_progress(self, val, text):
        self.progress_bar.setValue(val)
        self.lbl_status.setText(text)

    def on_finished(self, summary, transcript):
        self.current_summary = summary
        self.current_transcript = transcript

        combined_md = f"# Краткое содержание\n\n{summary}\n\n---\n\n# Полный транскрипт\n\n{transcript}"
        html = markdown.markdown(combined_md)
        self.text_viewer.setHtml(html)

        self.btn_start.setEnabled(True)
        self.btn_select.setEnabled(True)
        self.lbl_status.setText("Обработка завершена.")

        self.btn_export_md.setEnabled(True)
        self.btn_export_docx.setEnabled(True)
        self.btn_export_pdf.setEnabled(True)

    def on_error(self, err):
        QMessageBox.critical(self, "Ошибка", f"Произошла ошибка:\n{err}")
        self.btn_start.setEnabled(True)
        self.btn_select.setEnabled(True)
        self.lbl_status.setText("Ошибка при обработке.")

    def export_file(self, fmt):
        if not self.current_summary:
            return

        file_path, _ = QFileDialog.getSaveFileName(self, "Сохранить файл", f"Meeting_Result.{fmt}", f"*.{fmt}")
        if not file_path:
            return

        try:
            if fmt == "md":
                combined_md = f"# САММАРИ\n\n{self.current_summary}\n\n# ТРАНСКРИПТ\n\n{self.current_transcript}"
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(combined_md)

            elif fmt == "docx":
                doc = Document()
                doc.add_heading('Краткое содержание', 0)
                doc.add_paragraph(self.current_summary)
                doc.add_page_break()
                doc.add_heading('Полный транскрипт', 1)
                for line in self.current_transcript.split('\n\n'):
                    doc.add_paragraph(line.replace('**', ''))
                doc.save(file_path)

            elif fmt == "pdf":
                try:
                    self.text_viewer.printToPdf(file_path)
                    QMessageBox.information(self, "Успех", f"PDF файл успешно сохранен!")

                except Exception as e:
                    from PyQt6.QtPrintSupport import QPrinter
                    printer = QPrinter(QPrinter.PrinterMode.ScreenResolution)
                    printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
                    printer.setOutputFileName(file_path)
                    self.text_viewer.document().print(printer)

            QMessageBox.information(self, "Успех", f"Файл сохранен: {Path(file_path).name}")

        except Exception as e:
            QMessageBox.critical(self, "Ошибка сохранения", f"Не удалось сохранить файл:\n{str(e)}")


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = ProtonApp()
    window.show()
    sys.exit(app.exec())
